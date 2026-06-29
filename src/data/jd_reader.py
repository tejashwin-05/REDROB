"""Read job description from a .docx or plain text file."""

from __future__ import annotations

from pathlib import Path

from src.models.job_description import JobDescription
from src.utils.logging_utils import get_logger

logger = get_logger(__name__)


def read_job_description(path: str | Path) -> JobDescription:
    """
    Read a job description from a .docx or .txt file.
    Returns a JobDescription with the raw text.
    """
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"JD file not found: {path}")

    if path.suffix.lower() == ".docx":
        text = _read_docx(path)
    elif path.suffix.lower() in (".txt", ".md"):
        text = path.read_text(encoding="utf-8")
    else:
        raise ValueError(f"Unsupported JD format: {path.suffix}")

    logger.info(f"Loaded JD from {path.name} — {len(text)} chars")
    return JobDescription(raw_text=text, source_file=str(path))


def _read_docx(path: Path) -> str:
    """Extract all text paragraphs from a .docx file."""
    from docx import Document  # python-docx

    doc = Document(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Also grab table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt and txt not in paragraphs:
                    paragraphs.append(txt)

    return "\n".join(paragraphs)
